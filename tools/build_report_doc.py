# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "BaoCao_Cnpm_Bdtt_temp.docx"


def set_font(run, size=13, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def para(doc, text="", bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=10, name="Courier New")


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=12, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    para(doc)
    return table


def add_page_number_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Báo cáo bài thực hành - cnpm_bdtt")
    set_font(run, size=10)


def add_cover(doc):
    para(doc, "TRƯỜNG/ĐƠN VỊ ĐÀO TẠO", True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "KHOA CÔNG NGHỆ THÔNG TIN", True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "\nBÁO CÁO BÀI THỰC HÀNH", True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Đề tài: Hệ thống quét mã vạch và quản lý thuốc nhà thuốc", True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dự án: cnpm_bdtt", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "\nSinh viên thực hiện: ........................................", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Hình thức thực hiện: Cá nhân", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Ngày hoàn thành: 22/06/2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    para(doc, "MỤC LỤC", True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Ghi chú: Có thể cập nhật mục lục tự động trong Microsoft Word bằng References > Table of Contents trước khi in hoặc nộp bản cuối.")
    doc.add_page_break()


def chapter_1(doc):
    heading(doc, "Chương 1. Khảo sát và xác định yêu cầu")
    para(doc, "Chương này đặt nền tảng cho toàn bộ báo cáo bằng cách trình bày bối cảnh nghiệp vụ, nhu cầu sử dụng và những yêu cầu được rút ra từ chính hệ thống đã xây dựng. Đối với một bài thực hành phần mềm, phần khảo sát không chỉ nhằm mô tả vấn đề chung, mà còn phải làm rõ vì sao các chức năng trong source code là cần thiết và chúng phục vụ cho mục tiêu nào của người dùng.")

    heading(doc, "1.1. Giới thiệu bài toán", 2)
    para(doc, "Trong môi trường nhà thuốc, thông tin về thuốc thường gắn với nhiều yếu tố nghiệp vụ: mã vạch, tên thuốc, hoạt chất, nhà sản xuất, dạng bào chế, hàm lượng, hướng dẫn sử dụng, cảnh báo, số lô, hạn dùng và số lượng tồn kho. Nếu các thông tin này được tra cứu rời rạc hoặc ghi chép thủ công, nhân viên dễ mất thời gian khi phục vụ tại quầy, đặc biệt trong các tình huống cần kiểm tra nhanh thuốc còn hạn hay không, lô thuốc nào đang tồn, hoặc thuốc có cần lưu ý khi cấp phát hay không.")
    para(doc, "Dự án cnpm_bdtt được xây dựng như một hệ thống hỗ trợ nghiệp vụ cho nhà thuốc, trong đó ứng dụng Flutter đóng vai trò giao diện thao tác cho nhân viên, còn backend ASP.NET Core Web API chịu trách nhiệm xác thực, xử lý nghiệp vụ và làm việc với cơ sở dữ liệu SQL Server thông qua Entity Framework Core. Hệ thống hiện có các nhóm chức năng chính gồm đăng nhập, quét mã vạch, tra cứu thuốc, quản lý kho và lô thuốc, xác thực barcode/số lô, cảnh báo, báo cáo, quản lý nhà cung cấp và quản trị tài khoản nhân viên.")
    para(doc, "Điểm đáng chú ý của bài toán là các chức năng không tồn tại độc lập. Việc quét mã vạch liên quan trực tiếp đến danh mục thuốc và lịch sử quét; quản lý lô thuốc liên quan đến tồn kho, cảnh báo gần hết hạn và báo cáo; phân quyền người dùng quyết định ai được phép thêm, sửa hoặc xóa dữ liệu. Vì vậy, hệ thống được thiết kế theo hướng kết nối các nghiệp vụ thành một luồng vận hành tương đối hoàn chỉnh, phù hợp với phạm vi của một bài thực hành cá nhân.")

    heading(doc, "1.2. Product Vision", 2)
    para(doc, "Tầm nhìn của sản phẩm là xây dựng một công cụ nội bộ giúp nhà thuốc thao tác nhanh hơn, có dữ liệu tập trung hơn và giảm sai sót trong quá trình nhận diện, tra cứu và quản lý thuốc. Mã vạch được sử dụng như điểm bắt đầu của nhiều thao tác: từ một barcode, nhân viên có thể biết thuốc có tồn tại trong hệ thống hay không, xem thông tin thuốc, kiểm tra số lượng tồn và ghi nhận lịch sử quét.")
    para(doc, "Với người quản trị, hệ thống hướng đến việc cung cấp cái nhìn có cấu trúc về danh mục thuốc, lô thuốc, nhà cung cấp, tồn kho, cảnh báo và tài khoản nhân viên. Các dữ liệu này không chỉ phục vụ thao tác hằng ngày mà còn tạo cơ sở cho việc theo dõi tình trạng kho, phát hiện thuốc gần hết hạn hoặc tồn thấp, từ đó hỗ trợ quyết định nhập hàng và kiểm soát rủi ro.")
    para(doc, "Trong phạm vi hiện tại, sản phẩm chưa hướng đến triển khai thương mại hoặc thay thế các hệ thống quản lý nhà thuốc chuyên nghiệp. Thay vào đó, mục tiêu là xây dựng một nguyên mẫu có kiến trúc rõ ràng, có khả năng chạy demo, có phân quyền và có các luồng nghiệp vụ cốt lõi được hiện thực bằng code.")

    heading(doc, "1.3. Mục tiêu hệ thống", 2)
    para(doc, "Từ tầm nhìn trên, hệ thống được triển khai với các mục tiêu cụ thể sau:")
    for item in [
        "Cho phép người dùng đăng nhập bằng tài khoản nội bộ, nhận JWT và sử dụng hệ thống theo vai trò Admin hoặc Staff.",
        "Hỗ trợ nhân viên quét mã vạch thuốc, hiển thị kết quả tra cứu và lưu lại lịch sử quét phục vụ đối chiếu sau này.",
        "Cung cấp chức năng tìm kiếm thuốc theo tên, mã vạch, hoạt chất hoặc tên gọi khác; hiển thị thông tin cần thiết như giá bán, hướng dẫn sử dụng, cảnh báo và tồn kho.",
        "Quản lý lô thuốc gắn với hạn dùng, nhà cung cấp, số lượng tồn và ngưỡng cảnh báo tồn kho thấp.",
        "Ghi nhận các giao dịch nhập kho, xuất kho và điều chỉnh tồn kho để dữ liệu thay đổi có dấu vết.",
        "Tự động sinh cảnh báo đối với lô thuốc hết hạn, gần hết hạn trong vòng 90 ngày hoặc có tồn kho thấp.",
        "Cung cấp báo cáo tổng quan về số thuốc, số lô, tổng tồn kho, tồn kho thấp, lô hết hạn, lô gần hết hạn, lượt quét trong ngày và số lượng bán trong ngày.",
        "Cho phép Admin quản lý tài khoản Staff, bao gồm tạo, cập nhật, khóa/mở hoặc xóa tài khoản theo điều kiện dữ liệu liên quan.",
    ]:
        bullet(doc, item)

    heading(doc, "1.4. Phạm vi hệ thống", 2)
    para(doc, "Phạm vi cài đặt hiện tại gồm hai phần chính. Phần thứ nhất là ứng dụng Flutter nằm trong thư mục lib, được chia theo feature như auth, barcode_scan, medicine, inventory, alerts, reports, history, consultation, suppliers và admin. Phần thứ hai là backend ASP.NET Core Web API nằm trong thư mục backend, gồm các project Pharmacy.Api, Pharmacy.Application, Pharmacy.Domain và Pharmacy.Infrastructure.")
    para(doc, "Hệ thống phục vụ hai nhóm vai trò được định nghĩa trong enum UserRoleName là Admin và Staff. Staff có thể thực hiện các nghiệp vụ tại quầy như quét mã, tra cứu thuốc, xem lịch sử quét, xem cảnh báo, xác thực thuốc và xuất kho. Admin ngoài các chức năng trên còn có quyền quản lý thuốc, lô thuốc, nhà cung cấp, điều chỉnh kho và quản trị tài khoản nhân viên. Các quyền này được thể hiện ở cả backend thông qua thuộc tính Authorize(Roles = 'Admin') và frontend thông qua việc chỉ hiển thị menu quản trị khi token trả về có role Admin.")
    para(doc, "Những nội dung chưa thuộc phạm vi hiện tại gồm thanh toán, in hóa đơn, quản lý khách hàng, triển khai production, kết nối máy quét mã vạch chuyên dụng và bộ kiểm thử tự động hoàn chỉnh. Chức năng tư vấn thuốc đã có ConsultationService gọi Tavily và lọc kết quả từ tên miền Nhà thuốc Long Châu khi có API key, tuy nhiên báo cáo chỉ xem đây là chức năng tham khảo thông tin, không xem là hệ thống tư vấn y khoa hoàn chỉnh.")

    heading(doc, "1.5. Khảo sát hiện trạng", 2)
    para(doc, "Dựa trên dữ liệu seed trong PharmacyDbContext và DatabaseBootstrapper, hệ thống đang mô phỏng một nhà thuốc nhỏ với tài khoản demo, thuốc mẫu, lô thuốc mẫu, tồn kho và nhà cung cấp demo. Ba thuốc mẫu gồm Paracetamol 500mg, Ibuprofen 200mg và Amoxicillin 500mg; các barcode tương ứng được ghi trong README để phục vụ kiểm thử luồng quét mã. Cách tổ chức dữ liệu này cho phép người chấm có thể chạy hệ thống và kiểm tra nhanh các chức năng chính mà không cần nhập dữ liệu từ đầu.")
    para(doc, "Hiện trạng nghiệp vụ được mô phỏng trong project cho thấy một số vấn đề quan trọng: cần phân biệt quyền Admin và Staff; cần tìm thuốc nhanh bằng barcode; cần quản lý tồn kho theo từng lô thay vì chỉ theo tên thuốc; cần cảnh báo các lô gần hết hạn hoặc tồn thấp; và cần có báo cáo tổng quan để Admin nắm tình hình hệ thống. Các vấn đề này tương ứng trực tiếp với các controller và service đã được cài đặt trong backend.")

    heading(doc, "1.6. Stakeholder", 2)
    add_table(doc, ["Stakeholder", "Vai trò/Mối quan tâm"], [
        ["Chủ nhà thuốc/Admin", "Quan tâm đến dữ liệu thuốc, lô thuốc, tồn kho, cảnh báo, báo cáo và việc quản lý nhân viên."],
        ["Nhân viên nhà thuốc/Staff", "Cần thao tác nhanh tại quầy: đăng nhập, quét mã, tra cứu thuốc, xác thực lô, xuất kho và xem cảnh báo."],
        ["Người mua thuốc", "Gián tiếp hưởng lợi khi nhân viên cấp đúng thuốc, kiểm tra được hạn dùng và có thông tin cảnh báo cơ bản."],
        ["Giảng viên chấm bài", "Đánh giá khả năng phân tích yêu cầu, thiết kế hệ thống, hiện thực code, kiểm thử và trình bày báo cáo."],
    ])

    heading(doc, "1.7. Yêu cầu chức năng", 2)
    para(doc, "Các yêu cầu chức năng dưới đây được tổng hợp từ những endpoint, service và màn hình hiện có trong project. Do đó, chúng phản ánh phạm vi thật của hệ thống thay vì liệt kê các chức năng chưa được cài đặt.")
    add_table(doc, ["Mã", "Yêu cầu chức năng", "Căn cứ trong code"], [
        ["FR01", "Đăng nhập bằng username/password và nhận JWT.", "AuthController, AuthService, JwtTokenProvider"],
        ["FR02", "Quét barcode, trả kết quả thuốc và lưu lịch sử quét.", "ScansController, ScanService, BarcodeScanScreen"],
        ["FR03", "Tìm kiếm thuốc theo từ khóa, barcode, hoạt chất hoặc alias.", "MedicinesController, MedicineService.SearchAsync"],
        ["FR04", "Quản lý thuốc: tạo, cập nhật, xóa theo quyền Admin.", "MedicinesController, MedicineService"],
        ["FR05", "Quản lý lô thuốc, tồn kho và giao dịch nhập/xuất/điều chỉnh.", "InventoryController, InventoryService"],
        ["FR06", "Quản lý nhà cung cấp và ràng buộc không xóa khi đang dùng trong lô.", "SuppliersController, SupplierService"],
        ["FR07", "Xem lịch sử quét của người dùng hiện tại.", "ScanService.GetHistoryAsync"],
        ["FR08", "Xác thực thuốc bằng barcode và số lô; ghi log và cảnh báo khi không hợp lệ.", "VerificationController, VerificationService"],
        ["FR09", "Kiểm tra tương tác giữa các thuốc đã quét dựa trên MedicineInteraction.", "MedicineService.CheckInteractionsAsync"],
        ["FR10", "Sinh và hiển thị cảnh báo hết hạn, gần hết hạn và tồn kho thấp.", "AlertService.RefreshSystemAlertsAsync"],
        ["FR11", "Xem báo cáo tổng quan và danh sách thuốc/lô/lượt quét.", "ReportsController, ReportService"],
        ["FR12", "Quản trị tài khoản Staff: tạo, sửa, khóa/mở và xóa theo điều kiện.", "AdminController, AdminService"],
    ])

    heading(doc, "1.8. Yêu cầu phi chức năng", 2)
    para(doc, "Bên cạnh chức năng nghiệp vụ, hệ thống cũng cần đáp ứng một số yêu cầu phi chức năng để có thể vận hành ổn định trong môi trường demo và dễ bảo trì về sau.")
    for item in [
        "Bảo mật: backend dùng JWT Bearer Authentication; các controller nghiệp vụ đều yêu cầu Authorize, một số hành động quản trị giới hạn role Admin.",
        "Tính toàn vẹn dữ liệu: username và barcode được đặt unique index; InventoryItem gắn duy nhất với MedicineBatch; xóa thuốc hoặc nhà cung cấp có kiểm tra dữ liệu liên quan.",
        "Khả năng bảo trì: backend tách thành Api, Application, Domain và Infrastructure; frontend tách theo feature, giúp mỗi nhóm chức năng có model, service và screen riêng.",
        "Khả năng sử dụng: giao diện Flutter dùng NavigationDrawer để nhóm chức năng rõ ràng; màn hình quét mã dùng mobile_scanner để phù hợp với thao tác thực tế.",
        "Tính minh bạch nghiệp vụ: các thao tác kho được lưu thành InventoryTransaction; quét barcode được lưu thành ScanHistory; xác thực thuốc được lưu thành VerificationLog.",
        "Khả năng chạy demo: README cung cấp tài khoản demo, barcode demo và hướng dẫn chạy backend/frontend trong môi trường local.",
    ]:
        bullet(doc, item)


def chapter_2(doc):
    heading(doc, "Chương 2. Phân tích hệ thống theo Agile Scrum")
    para(doc, "Dự án được thực hiện cá nhân, vì vậy các nội dung Agile Scrum được sử dụng như một phương pháp tư duy và tổ chức công việc thay vì mô tả hoạt động của một nhóm phát triển. Các artifact như Product Backlog, Epic, User Story và Sprint Backlog được xây dựng dựa trên chức năng thật của hệ thống, nhằm thể hiện quá trình chuyển đổi từ nhu cầu nghiệp vụ thành các hạng mục phát triển cụ thể.")

    heading(doc, "2.1. Product Backlog", 2)
    para(doc, "Product Backlog được sắp xếp theo mức độ ưu tiên nghiệp vụ. Các hạng mục nền tảng như đăng nhập, phân quyền, quét barcode và tra cứu thuốc được đặt ở mức P1 vì chúng quyết định khả năng sử dụng hệ thống. Các chức năng báo cáo, cảnh báo, xác thực và quản trị được phát triển sau khi dữ liệu cốt lõi đã ổn định.")
    add_table(doc, ["Priority", "Backlog Item", "Mô tả", "Trạng thái"], [
        ["P1", "Đăng nhập và phân quyền", "Admin/Staff đăng nhập, nhận JWT, frontend hiển thị menu theo role.", "Đã có"],
        ["P1", "Quét mã vạch", "Dùng mobile_scanner, gọi /api/scans và lưu lịch sử quét.", "Đã có"],
        ["P1", "Tra cứu thuốc", "Tìm theo tên, barcode, hoạt chất, alias và xem thông tin chi tiết.", "Đã có"],
        ["P1", "Quản lý thuốc, lô và tồn kho", "Admin quản lý thuốc/lô; hệ thống lưu giao dịch kho.", "Đã có"],
        ["P2", "Cảnh báo nghiệp vụ", "Tự sinh cảnh báo hết hạn, gần hết hạn 90 ngày và tồn kho thấp.", "Đã có"],
        ["P2", "Xác thực barcode/số lô", "Đối chiếu barcode với số lô, ghi log và sinh cảnh báo khi không hợp lệ.", "Đã có"],
        ["P2", "Báo cáo", "Tổng hợp số thuốc, số lô, tồn kho, cảnh báo và lượt quét trong ngày.", "Đã có"],
        ["P3", "Tư vấn thuốc tham khảo", "Gọi Tavily, lọc nguồn Long Châu khi có API key cấu hình.", "Đã có service, phụ thuộc cấu hình"],
        ["P3", "Figma/Wireframe/Prototype", "Thiết kế giao diện trước khi code.", "Chưa có minh chứng trong project"],
    ])

    heading(doc, "2.2. Epic", 2)
    para(doc, "Các Epic được nhóm theo giá trị nghiệp vụ thay vì theo tên file. Cách nhóm này giúp nhìn thấy mối liên hệ giữa nhu cầu người dùng, API backend và màn hình frontend.")
    add_table(doc, ["Epic", "Ý nghĩa", "Chức năng liên quan"], [
        ["E1. Xác thực và phân quyền", "Bảo vệ dữ liệu nội bộ, phân biệt người quản trị và nhân viên.", "Login, JWT, Role Admin/Staff, NavigationDrawer theo role."],
        ["E2. Nhận diện và tra cứu thuốc", "Giúp nhân viên tìm thuốc nhanh tại quầy bằng barcode hoặc từ khóa.", "Scan, Medicine Search, Medicine Detail, Similar Medicines."],
        ["E3. Quản lý kho theo lô", "Theo dõi số lượng, hạn dùng, nhà cung cấp và lịch sử nhập/xuất.", "Batch, InventoryItem, InventoryTransaction, Supplier."],
        ["E4. Kiểm soát rủi ro nghiệp vụ", "Phát hiện lô hết hạn, gần hết hạn, tồn thấp hoặc xác thực không hợp lệ.", "AlertService, VerificationService, MedicineInteraction."],
        ["E5. Báo cáo và quản trị", "Cung cấp thông tin tổng hợp và hỗ trợ Admin quản lý nhân viên.", "ReportService, AdminService."],
    ])

    heading(doc, "2.3. User Story", 2)
    add_table(doc, ["Mã", "User Story"], [
        ["US01", "Là nhân viên, tôi muốn đăng nhập để hệ thống nhận diện tôi và cấp quyền sử dụng phù hợp."],
        ["US02", "Là nhân viên, tôi muốn quét barcode để nhanh chóng biết thuốc có trong cơ sở dữ liệu hay không."],
        ["US03", "Là nhân viên, tôi muốn xem thông tin thuốc sau khi quét để nắm giá bán, tồn kho, hạn dùng gần nhất và các cảnh báo cần lưu ý."],
        ["US04", "Là Admin, tôi muốn quản lý danh mục thuốc để dữ liệu tra cứu luôn được cập nhật."],
        ["US05", "Là Admin, tôi muốn quản lý lô thuốc để theo dõi hạn dùng, nhà cung cấp và số lượng tồn theo từng lô."],
        ["US06", "Là nhân viên, tôi muốn xuất kho khi bán thuốc để tồn kho phản ánh đúng tình trạng thực tế."],
        ["US07", "Là Admin, tôi muốn xem báo cáo để đánh giá tình hình thuốc, lô, tồn kho và hoạt động quét trong ngày."],
        ["US08", "Là Admin, tôi muốn quản lý tài khoản Staff để kiểm soát người được phép sử dụng hệ thống."],
    ])

    heading(doc, "2.4. Acceptance Criteria", 2)
    add_table(doc, ["User Story", "Acceptance Criteria"], [
        ["US01", "Đăng nhập đúng trả về token, họ tên, username và danh sách role; đăng nhập sai trả về Unauthorized."],
        ["US02", "Barcode tồn tại trả về thông tin thuốc; barcode không tồn tại trả về thông báo không tìm thấy nhưng vẫn ghi ScanHistory."],
        ["US03", "Thông tin thuốc hiển thị được tên, barcode, hoạt chất, nhà sản xuất, giá bán, tổng tồn, hạn dùng gần nhất và trạng thái kê đơn."],
        ["US04", "Chỉ Admin được tạo/sửa/xóa thuốc; barcode không được trùng; không xóa thuốc đã có lô/tồn kho."],
        ["US05", "Lô thuốc bắt buộc có thuốc, nhà cung cấp hợp lệ, số lô, ngày sản xuất, hạn dùng, số lượng và ngưỡng tồn thấp."],
        ["US06", "Xuất kho không cho phép tồn kho âm; mỗi thao tác xuất tạo InventoryTransaction loại Sale."],
        ["US07", "Báo cáo summary trả về số thuốc, số lô, tổng tồn, tồn thấp, hết hạn, gần hết hạn, lượt quét hôm nay và số lượng bán hôm nay."],
        ["US08", "Admin không được khóa tài khoản admin mặc định; tài khoản Staff có dữ liệu liên quan sẽ bị khóa thay vì xóa vật lý."],
    ])

    heading(doc, "2.5. Sprint Planning", 2)
    para(doc, "Kế hoạch sprint được xây dựng theo trình tự phụ thuộc kỹ thuật. Trước hết cần có domain model, database, đăng nhập và API cơ bản; sau đó mới có thể nối frontend với backend, phát triển kho, cảnh báo, báo cáo và quản trị. Cách chia này phù hợp với dự án cá nhân vì mỗi sprint tạo ra một phần có thể kiểm tra được.")
    add_table(doc, ["Sprint", "Mục tiêu", "Lý do sắp xếp"], [
        ["Sprint 1", "Khởi tạo kiến trúc backend/frontend, entity, DbContext, seed dữ liệu và đăng nhập.", "Tạo nền tảng để các chức năng sau có dữ liệu và bảo mật."],
        ["Sprint 2", "Xây dựng quét barcode, tra cứu thuốc, chi tiết thuốc và lịch sử quét.", "Đây là luồng nghiệp vụ trung tâm của ứng dụng."],
        ["Sprint 3", "Hoàn thiện quản lý lô, tồn kho, giao dịch kho và nhà cung cấp.", "Dữ liệu lô và tồn kho là cơ sở cho cảnh báo và báo cáo."],
        ["Sprint 4", "Bổ sung cảnh báo, xác thực, báo cáo, tư vấn tham khảo và quản trị nhân viên.", "Hoàn thiện các chức năng hỗ trợ vận hành và demo cuối kỳ."],
    ])

    heading(doc, "2.6. Sprint Backlog", 2)
    add_table(doc, ["Sprint", "Task tiêu biểu", "Kết quả trong source"], [
        ["Sprint 1", "Tạo solution backend nhiều project, cấu hình JWT, Swagger, CORS và database.", "Pharmacy.Api, Pharmacy.Application, Pharmacy.Domain, Pharmacy.Infrastructure."],
        ["Sprint 1", "Thiết kế entity và seed dữ liệu mẫu.", "Entities.cs, PharmacyDbContext.cs, DatabaseBootstrapper.cs."],
        ["Sprint 2", "Tạo API scan/search và giao diện quét mã.", "ScansController, MedicinesController, BarcodeScanScreen, MedicineSearchScreen."],
        ["Sprint 3", "Tạo nghiệp vụ kho và lô thuốc.", "InventoryController, InventoryService, SuppliersController, SupplierService."],
        ["Sprint 4", "Tạo cảnh báo, xác thực, báo cáo và admin.", "AlertService, VerificationService, ReportService, AdminService."],
    ])

    heading(doc, "2.7. Use Case Diagram", 2)
    para(doc, "Do báo cáo không có hình vẽ UML xuất từ công cụ chuyên dụng, sơ đồ use case được trình bày bằng mô tả văn bản để vẫn thể hiện đầy đủ tác nhân và chức năng.")
    code_block(doc, "Actor: Staff\n  -> Đăng nhập\n  -> Quét mã vạch\n  -> Tra cứu thuốc\n  -> Xem chi tiết thuốc\n  -> Xem lịch sử quét\n  -> Xác thực barcode/số lô\n  -> Xem cảnh báo\n  -> Xem báo cáo\n  -> Xuất kho\n\nActor: Admin kế thừa Staff\n  -> Quản lý thuốc\n  -> Quản lý lô thuốc\n  -> Nhập kho / điều chỉnh tồn kho\n  -> Quản lý nhà cung cấp\n  -> Quản trị tài khoản nhân viên")

    heading(doc, "2.8. Use Case Specification", 2)
    add_table(doc, ["Use Case", "Tác nhân", "Luồng chính", "Ngoại lệ"], [
        ["Đăng nhập", "Admin/Staff", "Nhập username/password, backend kiểm tra PasswordHash, tạo JWT và trả danh sách role.", "Sai thông tin hoặc tài khoản bị khóa thì trả Unauthorized."],
        ["Quét barcode", "Admin/Staff", "Ứng dụng nhận barcode từ camera, gọi /api/scans, backend tìm Medicine và lưu ScanHistory.", "Không tìm thấy thuốc thì trả Found=false."],
        ["Quản lý lô thuốc", "Admin", "Chọn thuốc, nhà cung cấp, nhập số lô, ngày sản xuất, hạn dùng, số lượng và ngưỡng tồn thấp.", "Thiếu nhà cung cấp hợp lệ hoặc số lượng âm thì báo lỗi."],
        ["Xuất kho", "Admin/Staff", "Chọn lô thuốc, nhập số lượng xuất, backend trừ tồn và lưu transaction Sale.", "Nếu tồn kho không đủ thì từ chối thao tác."],
        ["Xác thực thuốc", "Admin/Staff", "Nhập barcode và số lô, backend đối chiếu MedicineBatch và hạn dùng.", "Không khớp hoặc hết hạn thì ghi VerificationLog và tạo cảnh báo."],
        ["Quản trị nhân viên", "Admin", "Tạo/sửa/khóa/mở/xóa tài khoản Staff.", "Không được xóa hoặc khóa tài khoản admin mặc định."],
    ])

    heading(doc, "2.9. Activity Diagram", 2)
    code_block(doc, "Quy trình quét mã thuốc:\nBắt đầu -> Người dùng đăng nhập -> Mở màn hình Quét mã\n-> Camera đọc barcode -> Flutter kiểm tra barcode đã quét chưa\n-> Gửi POST /api/scans -> Backend tìm Medicine theo Barcode\n-> Lưu ScanHistory\n-> [Có thuốc] Trả thông tin thuốc, tồn kho, hạn dùng gần nhất\n-> [Không có thuốc] Trả thông báo không tìm thấy\n-> Flutter hiển thị kết quả -> Kết thúc")

    heading(doc, "2.10. Sequence Diagram", 2)
    code_block(doc, "Staff -> BarcodeScanScreen: Đưa mã vạch vào camera\nBarcodeScanScreen -> ScanService(Dart): scan(barcode)\nScanService(Dart) -> ApiClient: POST /api/scans + JWT\nApiClient -> ScansController: ScanRequest\nScansController -> ScanService(C#): ScanAsync(request, userId)\nScanService(C#) -> PharmacyDbContext: Tìm Medicine theo Barcode\nScanService(C#) -> PharmacyDbContext: Thêm ScanHistory\nPharmacyDbContext --> ScanService(C#): Lưu dữ liệu\nScanService(C#) --> ScansController: ScanResponse\nScansController --> Flutter App: JSON kết quả\nFlutter App --> Staff: Hiển thị thuốc hoặc thông báo không tìm thấy")

    heading(doc, "2.11. Class Diagram", 2)
    code_block(doc, "Entity\n  + Id, CreatedAt, UpdatedAt\n\nUser --< UserRole >-- Role\nMedicine --< MedicineAlias\nMedicine --< MedicineBatch -- InventoryItem\nMedicineBatch --< InventoryTransaction\nMedicine --< ScanHistory\nMedicine --< VerificationLog\nMedicineBatch --< Alert\nSupplier --< MedicineBatch\nMedicineInteraction --> MedicineA / MedicineB\nDispensingCheck --> User")


def chapter_3(doc):
    heading(doc, "Chương 3. Thiết kế hệ thống")
    para(doc, "Thiết kế hệ thống được xây dựng dựa trên kiến trúc đã hiện thực trong project, trong đó frontend và backend được tách rõ trách nhiệm. Frontend tập trung vào trải nghiệm thao tác, backend tập trung vào xác thực, phân quyền, xử lý nghiệp vụ và lưu trữ dữ liệu. Cách tổ chức này giúp hệ thống có thể mở rộng từng phần mà không làm toàn bộ code trở nên phụ thuộc chặt vào nhau.")

    heading(doc, "3.1. Kiến trúc hệ thống", 2)
    para(doc, "Hệ thống sử dụng mô hình client-server. Ứng dụng Flutter giao tiếp với backend qua HTTP/JSON. Lớp ApiClient trong Flutter chịu trách nhiệm tạo request, gắn token Bearer nếu có và xử lý lỗi trả về từ server. Ở phía backend, Program.cs cấu hình controller, Swagger, CORS, JWT Bearer Authentication, Authorization và đăng ký các service hạ tầng.")
    para(doc, "Backend được chia thành bốn project. Pharmacy.Domain chứa entity và enum nghiệp vụ; Pharmacy.Application chứa DTO và interface để định nghĩa hợp đồng nghiệp vụ; Pharmacy.Infrastructure chứa DbContext, service và logic làm việc với database; Pharmacy.Api chứa controller và cấu hình chạy ứng dụng. Đây là cách tổ chức tương đối rõ ràng, giúp giảm sự phụ thuộc trực tiếp giữa API layer và chi tiết lưu trữ dữ liệu.")
    code_block(doc, "Flutter App\n  | HTTP/JSON + JWT\nASP.NET Core Web API\n  | Controllers -> Application Interfaces -> Infrastructure Services\nEntity Framework Core\n  | DbContext / LINQ / SQL Server\nSQL Server Database")

    heading(doc, "3.2. Thiết kế giao diện bằng Figma", 2)
    para(doc, "")
    heading(doc, "3.3. Wireframe", 2)
    para(doc, "")
    heading(doc, "3.4. Prototype", 2)
    para(doc, "")

    heading(doc, "3.5. Thiết kế cơ sở dữ liệu", 2)
    para(doc, "Cơ sở dữ liệu được thiết kế theo hướng phục vụ nghiệp vụ nhà thuốc. Nhóm User, Role và UserRole đảm nhiệm xác thực và phân quyền. Nhóm Medicine, MedicineAlias, MedicineInteraction lưu dữ liệu danh mục thuốc và quan hệ tương tác. Nhóm MedicineBatch, InventoryItem, InventoryTransaction và Supplier phục vụ quản lý lô, tồn kho, giao dịch kho và nguồn cung. Các bảng ScanHistory, VerificationLog và Alert giúp hệ thống lưu dấu vết thao tác và cảnh báo rủi ro.")
    para(doc, "Một số ràng buộc quan trọng đã được thể hiện trong DbContext: Username là duy nhất, Barcode của thuốc là duy nhất, InventoryItem gắn duy nhất với MedicineBatch, và MedicineInteraction dùng quan hệ Restrict để tránh xóa dây chuyền không mong muốn giữa hai thuốc có tương tác.")

    heading(doc, "3.6. ERD", 2)
    code_block(doc, "Users (1) -- (N) UserRoles (N) -- (1) Roles\nMedicines (1) -- (N) MedicineAliases\nMedicines (1) -- (N) MedicineBatches\nSuppliers (1) -- (N) MedicineBatches\nMedicineBatches (1) -- (1) InventoryItems\nMedicineBatches (1) -- (N) InventoryTransactions\nUsers (1) -- (N) InventoryTransactions\nUsers (1) -- (N) ScanHistories\nMedicines (1) -- (N) ScanHistories\nMedicines (1) -- (N) VerificationLogs\nMedicineBatches (1) -- (N) Alerts")

    heading(doc, "3.7. Database Schema", 2)
    add_table(doc, ["Bảng", "Trường chính", "Vai trò trong hệ thống"], [
        ["Users", "Id, FullName, Username, PasswordHash, IsActive", "Lưu tài khoản người dùng và trạng thái hoạt động."],
        ["Roles", "Id, Name", "Lưu vai trò Admin hoặc Staff."],
        ["UserRole", "UserId, RoleId", "Liên kết nhiều-nhiều giữa người dùng và vai trò."],
        ["Medicines", "Id, Name, Barcode, ActiveIngredient, Manufacturer, DosageForm, Strength, SalePrice", "Lưu danh mục thuốc và thông tin tra cứu."],
        ["MedicineAliases", "Id, MedicineId, Alias", "Lưu tên gọi khác để mở rộng khả năng tìm kiếm."],
        ["MedicineInteractions", "MedicineAId, MedicineBId, Severity, Description", "Lưu cảnh báo tương tác giữa hai thuốc."],
        ["Suppliers", "Id, Name, Phone, Address", "Lưu nhà cung cấp cho các lô thuốc."],
        ["MedicineBatches", "Id, MedicineId, SupplierId, BatchNumber, ManufactureDate, ExpiryDate, InitialQuantity", "Lưu thông tin lô thuốc và hạn dùng."],
        ["InventoryItems", "Id, MedicineBatchId, Quantity, LowStockThreshold", "Lưu số lượng tồn hiện tại theo từng lô."],
        ["InventoryTransactions", "Id, MedicineBatchId, UserId, Type, Quantity, Note", "Ghi nhận giao dịch nhập, bán/xuất và điều chỉnh kho."],
        ["ScanHistories", "Id, MedicineId, UserId, Barcode, Found", "Ghi lịch sử quét barcode của người dùng."],
        ["VerificationLogs", "Id, MedicineId, UserId, Barcode, BatchNumber, IsVerified, Message", "Ghi lịch sử xác thực barcode và số lô."],
        ["Alerts", "Id, MedicineId, MedicineBatchId, Type, Severity, Title, Message, IsResolved", "Lưu cảnh báo hết hạn, gần hết hạn, tồn thấp hoặc xác thực."],
    ])

    heading(doc, "3.8. Deployment Diagram", 2)
    para(doc, "Trong phạm vi bài thực hành, hệ thống được triển khai ở môi trường local. Backend chạy trên máy tính phát triển, kết nối SQL Server và mở Swagger để kiểm thử API. Ứng dụng Flutter có thể chạy trên Android Emulator, khi đó base URL mặc định dùng địa chỉ 10.0.2.2 để truy cập localhost của máy tính. Nếu chạy trên thiết bị thật, README hướng dẫn truyền API_BASE_URL bằng IP của máy tính.")
    code_block(doc, "Máy tính phát triển\n  - ASP.NET Core Web API: http://localhost:5000\n  - Swagger: http://localhost:5000/swagger\n  - SQL Server: PharmacyBarcodeDb\n\nThiết bị chạy Flutter\n  - Android Emulator: http://10.0.2.2:5000/api\n  - Thiết bị thật: http://<IP_MAY_TINH>:5000/api")


def chapter_4(doc):
    heading(doc, "Chương 4. Xây dựng hệ thống")
    para(doc, "Chương này trình bày quá trình hiện thực hệ thống dựa trên source code. Nội dung được tổ chức theo các lớp của hệ thống: công nghệ sử dụng, cấu trúc source, frontend, backend, API, database và các chức năng chính. Mục tiêu của chương không chỉ là kể tên công nghệ, mà còn làm rõ mỗi công nghệ đang đảm nhiệm vai trò gì trong sản phẩm.")

    heading(doc, "4.1. Công nghệ sử dụng", 2)
    add_table(doc, ["Thành phần", "Công nghệ", "Vai trò"], [
        ["Frontend", "Flutter, Dart, Material UI", "Xây dựng giao diện đa nền tảng và điều hướng chức năng."],
        ["Quét mã", "mobile_scanner", "Đọc barcode bằng camera trên màn hình quét mã."],
        ["Giao tiếp API", "http, ApiClient", "Gửi request HTTP/JSON và xử lý token/lỗi."],
        ["Lưu token", "shared_preferences", "Lưu JWT ở phía ứng dụng Flutter."],
        ["Backend", "ASP.NET Core Web API .NET 8, C#", "Xây dựng REST API và xử lý nghiệp vụ."],
        ["Database", "SQL Server, Entity Framework Core 8", "Lưu trữ dữ liệu và truy vấn bằng DbContext/LINQ."],
        ["Bảo mật", "JWT Bearer Authentication", "Xác thực request và phân quyền theo role."],
        ["Tài liệu API", "Swagger/Swashbuckle", "Hỗ trợ kiểm thử API trong môi trường development."],
    ])

    heading(doc, "4.2. Cấu trúc source code", 2)
    para(doc, "Source code được chia thành hai phần lớn. Thư mục lib chứa ứng dụng Flutter, tổ chức theo feature để mỗi chức năng có screen, model và service riêng. Thư mục backend chứa solution ASP.NET Core, được chia thành nhiều project để tách domain model, DTO/interface, service hạ tầng và controller.")
    code_block(doc, "cnpm_bdtt/\n  lib/\n    app/                 Khung ứng dụng và ShellScreen\n    core/                ApiClient, config, token storage, tiện ích chung\n    features/            auth, barcode_scan, medicine, inventory, alerts, reports, history, consultation, suppliers, admin\n  backend/\n    Pharmacy.Api/        Controllers, Program.cs, JWT, Swagger\n    Pharmacy.Application/DTOs và interface nghiệp vụ\n    Pharmacy.Domain/     Entity và enum\n    Pharmacy.Infrastructure/DbContext, service, seed dữ liệu")

    heading(doc, "4.3. Frontend", 2)
    para(doc, "Frontend bắt đầu từ main.dart, gọi PharmacyApp và sau khi đăng nhập thành công sẽ đi vào ShellScreen. ShellScreen tạo danh sách màn hình theo role: các chức năng chung như Quét mã, Tư vấn, Tìm thuốc, Kho và lô, Lịch sử quét, Xác thực, Cảnh báo và Danh sách thuốc hiển thị cho người dùng đã đăng nhập; riêng Nhà cung ứng và Quản trị chỉ hiển thị khi danh sách role có Admin.")
    para(doc, "Màn hình BarcodeScanScreen là một ví dụ thể hiện rõ sự kết hợp giữa giao diện và nghiệp vụ. Camera đọc barcode bằng MobileScanner, sau đó màn hình gọi ScanService phía Flutter, service này tiếp tục gọi API /api/scans thông qua ApiClient. Kết quả trả về được đưa vào danh sách hiển thị; nếu thuốc tồn tại, người dùng có thể mở màn hình chi tiết thuốc để xem thêm thông tin và thuốc tương tự.")
    para(doc, "ApiClient đóng vai trò lớp trung gian quan trọng. Mọi request GET, POST, PUT, PATCH, DELETE đều đi qua lớp này; token được đọc từ TokenStorage và gắn vào header Authorization. Khi backend trả 401, ApiClient xóa token và gọi onUnauthorized để đưa người dùng về trạng thái cần đăng nhập lại. Nhờ đó, xử lý lỗi xác thực không bị lặp lại ở từng màn hình.")

    heading(doc, "4.4. Backend", 2)
    para(doc, "Backend được xây dựng theo hướng controller mỏng và service xử lý nghiệp vụ. Controller nhận request, lấy userId từ ClaimsPrincipal khi cần, kiểm tra role bằng Authorize và gọi service tương ứng. Ví dụ, ScansController gọi IScanService, InventoryController gọi IInventoryService, ReportsController gọi IReportService. Điều này giúp controller tập trung vào HTTP contract, trong khi nghiệp vụ nằm ở Infrastructure service.")
    para(doc, "Program.cs cấu hình các thành phần nền tảng: controller, Swagger, CORS cho mobile development, JWT Bearer Authentication, Authorization và Dependency Injection. Khi ứng dụng khởi động, DatabaseBootstrapper.EnsureSchemaAndDefaultsAsync được gọi để đảm bảo database có schema và dữ liệu mặc định. Đây là lựa chọn phù hợp với môi trường demo vì người chấm có thể chạy backend mà không cần chuẩn bị migration phức tạp.")
    para(doc, "Các service hiện thực nghiệp vụ có ràng buộc khá rõ. MedicineService kiểm tra barcode trùng và không cho xóa thuốc đã có lô. InventoryService không cho tồn kho âm, tạo transaction khi nhập/xuất/điều chỉnh và làm mới cảnh báo sau mỗi thay đổi. AlertService xóa cảnh báo hệ thống cũ rồi tính lại cảnh báo hết hạn, gần hết hạn trong 90 ngày và tồn kho thấp. AdminService bảo vệ tài khoản admin mặc định và khóa Staff thay vì xóa nếu tài khoản đã có dữ liệu liên quan.")

    heading(doc, "4.5. API", 2)
    add_table(doc, ["Nhóm API", "Endpoint tiêu biểu", "Ý nghĩa nghiệp vụ"], [
        ["Auth", "POST /api/auth/login", "Đăng nhập, kiểm tra mật khẩu và cấp JWT."],
        ["Medicines", "GET /api/medicines/search; GET /api/medicines/barcode/{barcode}", "Tra cứu thuốc và lấy thông tin theo barcode."],
        ["Scans", "POST /api/scans; GET /api/scans/history", "Quét barcode và xem lịch sử quét của người dùng."],
        ["Inventory", "GET /api/inventory/batches; POST /api/inventory/import; POST /api/inventory/export", "Quản lý lô và giao dịch kho."],
        ["Verification", "POST /api/verification", "Đối chiếu barcode với số lô, ghi log xác thực."],
        ["Alerts", "GET /api/alerts", "Xem các cảnh báo chưa được xử lý."],
        ["Reports", "GET /api/reports/summary", "Xem số liệu tổng quan phục vụ quản trị."],
        ["Suppliers", "GET/POST/PUT/DELETE /api/suppliers", "Quản lý nhà cung cấp, giới hạn thao tác ghi cho Admin."],
        ["Admin", "GET /api/admin/users; POST /api/admin/staff", "Quản trị tài khoản Staff."],
        ["Consultation", "POST /api/consultation/medicine", "Tra cứu thông tin tham khảo qua Tavily khi có cấu hình API key."],
    ])

    heading(doc, "4.6. Database", 2)
    para(doc, "Database dùng SQL Server và được truy cập qua PharmacyDbContext. DbContext khai báo các DbSet tương ứng với entity nghiệp vụ, đồng thời cấu hình index, quan hệ và dữ liệu seed. Dữ liệu seed gồm role Admin/Staff, tài khoản admin và staff, nhà cung cấp demo, ba thuốc mẫu, alias, tương tác thuốc, lô thuốc và số lượng tồn ban đầu.")
    para(doc, "Việc sử dụng EnsureCreated giúp giảm công sức thiết lập khi chạy bài thực hành. Tuy nhiên, README backend cũng ghi chú rằng khi cần nộp bản hoàn chỉnh hoặc phát triển lâu dài, nên chuyển sang EF Core Migration để quản lý lịch sử thay đổi schema tốt hơn. Nhận xét này được giữ trong báo cáo như một hạn chế và hướng phát triển hợp lý.")

    heading(doc, "4.7. Các chức năng chính", 2)
    for item in [
        "Đăng nhập và phân quyền bằng JWT, phân biệt Admin và Staff ở cả backend lẫn frontend.",
        "Quét mã vạch bằng camera, lưu ScanHistory và hiển thị kết quả thuốc hoặc thông báo không tìm thấy.",
        "Tra cứu thuốc, xem thông tin chi tiết, giá bán, tổng tồn, hạn dùng gần nhất, cảnh báo và thuốc tương tự.",
        "Quản lý thuốc, lô thuốc, nhà cung cấp, tồn kho và giao dịch kho.",
        "Tự động cảnh báo thuốc hết hạn, gần hết hạn trong 90 ngày và tồn kho thấp.",
        "Xác thực thuốc theo barcode và số lô, ghi VerificationLog và tạo cảnh báo khi không hợp lệ.",
        "Báo cáo tổng quan và báo cáo danh sách thuốc, lô, lô gần hết hạn, lô hết hạn, lượt quét trong ngày.",
        "Quản trị tài khoản nhân viên với cơ chế bảo vệ tài khoản admin mặc định.",
        "Tư vấn thuốc tham khảo qua Tavily/Long Châu khi có cấu hình API key.",
    ]:
        bullet(doc, item)

    heading(doc, "4.8. Kết quả từng Sprint", 2)
    add_table(doc, ["Sprint", "Kết quả đạt được"], [
        ["Sprint 1", "Hoàn thành kiến trúc backend/frontend, cấu hình JWT, Swagger, database và dữ liệu demo."],
        ["Sprint 2", "Hoàn thành quét barcode, tra cứu thuốc, chi tiết thuốc và lịch sử quét."],
        ["Sprint 3", "Hoàn thành quản lý kho, lô thuốc, giao dịch kho và nhà cung cấp."],
        ["Sprint 4", "Hoàn thành cảnh báo, xác thực, báo cáo, tư vấn tham khảo và quản trị nhân viên."],
    ])

    heading(doc, "4.9. Hình ảnh minh họa", 2)
    para(doc, "")


def chapter_5(doc):
    heading(doc, "Chương 5. Kiểm thử")
    para(doc, "Kiểm thử trong báo cáo được trình bày theo hướng kiểm thử thủ công có căn cứ từ code và dữ liệu demo. Project hiện chưa có bộ unit test hoặc integration test đầy đủ, vì vậy báo cáo không ghi nhận sai rằng các kiểm thử tự động đã được triển khai. Thay vào đó, các test case dưới đây mô tả những tình huống có thể kiểm tra trực tiếp qua Swagger và giao diện Flutter.")

    heading(doc, "5.1. Kế hoạch kiểm thử", 2)
    para(doc, "Mục tiêu kiểm thử là xác nhận các luồng nghiệp vụ cốt lõi hoạt động đúng: đăng nhập, phân quyền, quét barcode, tra cứu thuốc, quản lý kho, xác thực thuốc, cảnh báo, báo cáo và quản trị nhân viên. Môi trường kiểm thử gồm backend chạy tại http://localhost:5000, Swagger tại /swagger và app Flutter kết nối API qua http://10.0.2.2:5000/api khi chạy Android Emulator.")
    para(doc, "Dữ liệu kiểm thử dựa trên seed data có sẵn trong project: tài khoản admin/admin123, staff/staff123 và các barcode demo được ghi trong README. Cách này giúp quá trình kiểm thử có thể lặp lại và không phụ thuộc vào việc nhập dữ liệu thủ công trước khi chạy.")

    heading(doc, "5.2. Test Case", 2)
    add_table(doc, ["TC", "Chức năng", "Dữ liệu", "Kết quả mong đợi"], [
        ["TC01", "Đăng nhập đúng", "admin/admin123", "Trả JWT, role Admin và hiển thị menu quản trị."],
        ["TC02", "Đăng nhập sai", "admin/sai_mat_khau", "Trả Unauthorized và hiển thị lỗi đăng nhập."],
        ["TC03", "Đăng nhập Staff", "staff/staff123", "Không hiển thị menu Nhà cung ứng và Quản trị."],
        ["TC04", "Quét barcode tồn tại", "8938505974190", "Hiển thị Paracetamol 500mg và lưu ScanHistory."],
        ["TC05", "Quét barcode không tồn tại", "0000000000000", "Hiển thị không tìm thấy, Found=false và vẫn lưu lịch sử."],
        ["TC06", "Tìm thuốc", "para", "Danh sách trả về thuốc phù hợp theo tên/hoạt chất/barcode/alias."],
        ["TC07", "Xuất kho hợp lệ", "Chọn lô còn tồn, số lượng nhỏ hơn tồn", "Tồn kho giảm, tạo InventoryTransaction loại Sale."],
        ["TC08", "Xuất kho vượt tồn", "Số lượng lớn hơn tồn", "Backend từ chối với lỗi tồn kho không đủ."],
        ["TC09", "Cảnh báo tồn thấp", "Lô có Quantity <= LowStockThreshold", "AlertService sinh cảnh báo LowStock."],
        ["TC10", "Xác thực đúng", "Barcode và số lô khớp, lô chưa hết hạn", "IsVerified=true, severity Info."],
        ["TC11", "Xác thực sai", "Barcode/số lô không khớp hoặc lô hết hạn", "IsVerified=false, ghi VerificationLog và tạo Alert."],
        ["TC12", "Xem báo cáo", "GET /api/reports/summary", "Trả số liệu MedicineCount, BatchCount, TotalInventoryQuantity và các chỉ số liên quan."],
    ])

    heading(doc, "5.3. Test Scenario", 2)
    for item in [
        "Admin đăng nhập, tạo hoặc cập nhật dữ liệu thuốc, tạo lô, nhập kho, xem cảnh báo và báo cáo tổng quan.",
        "Staff đăng nhập, quét barcode thuốc, xem chi tiết thuốc, xuất kho khi bán và kiểm tra lại lịch sử quét.",
        "Người dùng nhập barcode và số lô để xác thực thuốc; hệ thống ghi log và cảnh báo nếu dữ liệu không hợp lệ.",
        "Admin tạo tài khoản Staff mới, cập nhật thông tin nhân viên, khóa/mở tài khoản và kiểm tra sự thay đổi ở giao diện.",
        "Hệ thống làm mới cảnh báo sau khi thay đổi tồn kho hoặc lô thuốc, bảo đảm danh sách cảnh báo phản ánh dữ liệu hiện tại.",
    ]:
        number(doc, item)

    heading(doc, "5.4. Unit Test", 2)
    para(doc, "")
    heading(doc, "5.5. Integration Test", 2)
    para(doc, "")
    heading(doc, "5.6. System Test", 2)
    para(doc, "")

    heading(doc, "5.7. Kết quả kiểm thử", 2)
    para(doc, "Kết quả kiểm thử thủ công dự kiến cho thấy các luồng chính có thể được xác nhận bằng dữ liệu demo và Swagger. Những chức năng có ràng buộc nghiệp vụ như không cho tồn kho âm, không xóa thuốc đã có lô, không khóa admin mặc định và tự sinh cảnh báo đều có logic tương ứng trong service backend.")
    para(doc, "Tuy nhiên, vì source hiện chưa có minh chứng về unit test, integration test hoặc system test tự động, các mục này được giữ trống. Đây là hạn chế cần được bổ sung nếu dự án tiếp tục phát triển sau phạm vi bài thực hành.")


def chapter_6(doc):
    heading(doc, "Chương 6. Triển khai và đánh giá")
    para(doc, "Chương cuối tổng hợp môi trường chạy hệ thống, kết quả đạt được, hạn chế và hướng phát triển. Do dự án được thực hiện cá nhân, các phần Sprint Review và Sprint Retrospective được trình bày dưới góc độ tự đánh giá quá trình phát triển, không mô tả hoạt động trao đổi nhóm.")

    heading(doc, "6.1. Môi trường triển khai", 2)
    for item in [
        "Hệ điều hành: Windows.",
        "Backend: .NET 8 SDK, ASP.NET Core Web API, chạy bằng lệnh dotnet run --project Pharmacy.Api.",
        "Database: SQL Server Express hoặc SQL Server Developer, database PharmacyBarcodeDb.",
        "Frontend: Flutter SDK, Android Emulator hoặc thiết bị thật.",
        "API document: Swagger tại http://localhost:5000/swagger trong môi trường Development.",
        "Khi chạy Android Emulator, app gọi API qua http://10.0.2.2:5000/api; khi chạy thiết bị thật cần truyền IP máy tính qua API_BASE_URL.",
    ]:
        bullet(doc, item)

    heading(doc, "6.2. Sprint Review", 2)
    para(doc, "Sau mỗi sprint, kết quả được rà soát dựa trên khả năng chạy được của chức năng. Sprint đầu tiên tạo nền tảng dự án và dữ liệu demo. Sprint thứ hai chứng minh được luồng quét mã và tra cứu thuốc. Sprint thứ ba mở rộng hệ thống sang quản lý kho theo lô, giúp dữ liệu có chiều sâu nghiệp vụ hơn. Sprint cuối cùng hoàn thiện các chức năng hỗ trợ quản trị như cảnh báo, báo cáo, xác thực và quản lý nhân viên.")
    para(doc, "Nhìn chung, sản phẩm sau các sprint đã đạt được mục tiêu của một bài thực hành: có backend chạy độc lập, có app Flutter kết nối API, có dữ liệu demo, có phân quyền và có các luồng nghiệp vụ liên kết với nhau.")

    heading(doc, "6.3. Sprint Retrospective", 2)
    para(doc, "Quá trình thực hiện cho thấy việc chia code theo lớp và theo feature giúp dự án dễ kiểm soát hơn. Khi backend có DTO, interface và service rõ ràng, frontend có thể gọi API qua một lớp ApiClient thống nhất. Điều này giảm lặp code và giúp các màn hình không phải tự xử lý chi tiết xác thực hoặc lỗi HTTP.")
    para(doc, "Bên cạnh đó, dự án cũng cho thấy một số điểm cần cải thiện. Việc dùng EnsureCreated thuận tiện cho demo nhưng chưa phù hợp với quản lý schema lâu dài. Một số nội dung như Figma, prototype và test tự động chưa được bổ sung. Ngoài ra, chức năng tư vấn thuốc phụ thuộc API key Tavily nên cần cấu hình riêng và cần được kiểm soát kỹ nếu dùng trong ngữ cảnh y tế thực tế.")

    heading(doc, "6.4. Kết quả đạt được", 2)
    for item in [
        "Xây dựng được ứng dụng Flutter có cấu trúc theo feature và giao diện điều hướng rõ ràng.",
        "Xây dựng được backend ASP.NET Core Web API .NET 8 với JWT, Swagger, CORS và phân quyền Admin/Staff.",
        "Thiết kế được domain model cho thuốc, alias, tương tác thuốc, lô thuốc, tồn kho, giao dịch kho, lịch sử quét, xác thực, cảnh báo và người dùng.",
        "Hiện thực được các nghiệp vụ chính: đăng nhập, quét barcode, tra cứu thuốc, quản lý kho/lô, nhà cung cấp, cảnh báo, xác thực, báo cáo và quản trị nhân viên.",
        "Có dữ liệu demo, tài khoản demo và hướng dẫn chạy local trong README.",
        "Có cấu trúc source đủ rõ để tiếp tục mở rộng bằng migration, test tự động hoặc triển khai server thật.",
    ]:
        bullet(doc, item)

    heading(doc, "6.5. Hạn chế", 2)
    for item in [
        "Chưa có file thiết kế Figma, wireframe và prototype để chứng minh quá trình thiết kế giao diện trước khi code.",
        "Chưa có bộ unit test, integration test và system test tự động trong source.",
        "Database hiện dùng EnsureCreated và bootstrapper, phù hợp demo nhưng chưa tối ưu cho quản lý phiên bản schema dài hạn.",
        "Chưa triển khai production, chưa cấu hình HTTPS/deployment server và chưa có quy trình backup dữ liệu.",
        "Chức năng tư vấn thuốc phụ thuộc Tavily API key và nguồn Long Châu, chỉ nên xem là tham khảo, không thay thế tư vấn của dược sĩ hoặc bác sĩ.",
        "Một số thông báo trong code backend/frontend còn có chuỗi không dấu hoặc chưa được chuẩn hóa hoàn toàn cho trải nghiệm người dùng cuối.",
    ]:
        bullet(doc, item)

    heading(doc, "6.6. Hướng phát triển", 2)
    for item in [
        "Bổ sung Figma, wireframe, prototype và ảnh minh họa màn hình thật để hoàn thiện hồ sơ thiết kế.",
        "Chuyển database sang EF Core Migration, xây dựng dữ liệu mẫu bằng seed có kiểm soát và chuẩn hóa quy trình update database.",
        "Viết unit test cho service nghiệp vụ, integration test cho API và một số test widget cho Flutter.",
        "Bổ sung chức năng bán hàng, in hóa đơn, quản lý khách hàng và lịch sử mua thuốc.",
        "Triển khai backend lên server nội bộ hoặc cloud, cấu hình HTTPS, biến môi trường và cơ chế backup.",
        "Hoàn thiện chức năng tư vấn thuốc bằng nguồn dữ liệu đáng tin cậy, có cảnh báo pháp lý và giới hạn rõ nội dung tham khảo.",
        "Tích hợp máy quét barcode chuyên dụng hoặc tối ưu trải nghiệm camera cho môi trường nhà thuốc thực tế.",
    ]:
        bullet(doc, item)


def build():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    add_page_number_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
